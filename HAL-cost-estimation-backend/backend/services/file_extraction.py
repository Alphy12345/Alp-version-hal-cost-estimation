import os
import re
import shutil
import sys
from typing import Optional
from fastapi import UploadFile

# Check if Tesseract is in PATH
def check_tesseract_installed():
    """Check if tesseract command is available"""
    tesseract_cmd = 'tesseract'
    if sys.platform == 'win32':
        tesseract_cmd = 'tesseract.exe'
    
    for path in os.environ.get('PATH', '').split(os.pathsep):
        exe_path = os.path.join(path, tesseract_cmd)
        if os.path.exists(exe_path):
            print(f"Found Tesseract at: {exe_path}")
            return True
    return False

TESSERACT_IN_PATH = check_tesseract_installed()
print(f"Tesseract in PATH: {TESSERACT_IN_PATH}")

try:
    import pytesseract
    from pdf2image import convert_from_path
    from PIL import Image
    
    # Set explicit path to Tesseract on Windows
    if sys.platform == 'win32':
        possible_tesseract_paths = [
            r'C:\Program Files\Tesseract-OCR\tesseract.exe',
            r'C:\Program Files (x86)\Tesseract-OCR\tesseract.exe',
            os.path.expandvars(r'%LOCALAPPDATA%\Programs\Tesseract-OCR\tesseract.exe'),
        ]
        for tesseract_path in possible_tesseract_paths:
            if os.path.exists(tesseract_path):
                pytesseract.pytesseract.tesseract_cmd = tesseract_path
                print(f"Set Tesseract path: {tesseract_path}")
                break
        
        # Set Poppler path for pdf2image on Windows
        possible_poppler_paths = [
            r'C:\poppler\poppler-24.08.0\Library\bin',
            r'C:\poppler\Library\bin',
            r'C:\poppler\bin',
        ]
        for poppler_path in possible_poppler_paths:
            if os.path.exists(poppler_path):
                os.environ['PATH'] = poppler_path + os.pathsep + os.environ.get('PATH', '')
                print(f"Added Poppler to PATH: {poppler_path}")
                break
    
    TESSERACT_AVAILABLE = True
    try:
        version = pytesseract.get_tesseract_version()
        print(f"Tesseract version: {version}")
    except Exception as e:
        print(f"Tesseract version check failed: {e}")
        TESSERACT_AVAILABLE = False
except Exception as e:
    TESSERACT_AVAILABLE = False
    print(f"WARNING: OCR libraries not available: {e}")

try:
    import pytesseract
    from pdf2image import convert_from_path
    from PIL import Image
    TESSERACT_AVAILABLE = True
    print("Tesseract OCR available")
except ImportError as e:
    TESSERACT_AVAILABLE = False
    print(f"WARNING: OCR libraries not available: {e}")

try:
    import pdfplumber
    print("PDFPlumber available for text-based PDF extraction")
except ImportError as e:
    pdfplumber = None
    print(f"WARNING: pdfplumber not available: {e}")

try:
    import docx
    print("python-docx available for Word document extraction")
except ImportError as e:
    docx = None
    print(f"WARNING: python-docx not available: {e}")

try:
    import openpyxl
    print("openpyxl available for Excel .xlsx extraction")
except ImportError as e:
    openpyxl = None
    print(f"WARNING: openpyxl not available: {e}")

try:
    import xlrd
    print("xlrd available for Excel .xls extraction")
except ImportError as e:
    xlrd = None
    print(f"WARNING: xlrd not available: {e}")

class FileExtractionService:
    """Service to extract operation details from uploaded files."""
    
    SUPPORTED_EXTENSIONS = {'.pdf', '.png', '.jpg', '.jpeg', '.tiff', '.bmp', '.txt', '.csv', '.doc', '.docx', '.xls', '.xlsx'}
    
    # Keywords to search for in extracted text - expanded
    OPERATION_KEYWORDS = ['operation', 'operation type', 'op type', 'process', 'op:', 'operation:', 'opn', 'op.']
    MATERIAL_KEYWORDS = ['material', 'matl', 'mat:', 'material:', 'matl.', 'materials']
    MACHINE_KEYWORDS = ['machine', 'equipment', 'machine name', 'mach:', 'machine:', 'mach', 'mach.']
    MAN_HOURS_KEYWORDS = ['man hours', 'manhours', 'man hrs', 'hours per unit', 'hrs/unit', 'hrs', 'hours', 'man hr', 'hr/unit']
    DUTY_KEYWORDS = ['duty', 'duty category', 'load', 'duty:', 'duty cat', 'category']
    DIAMETER_KEYWORDS = ['diameter', 'dia', 'diam', 'd:', 'dia:', 'dia.', 'Ø', 'phi']
    LENGTH_KEYWORDS = ['length', 'len', 'l', 'len:', 'length:', 'l.']
    BREADTH_KEYWORDS = ['breadth', 'width', 'b', 'w', 'breadth:', 'width:', 'b.', 'w.']
    HEIGHT_KEYWORDS = ['height', 'h', 'thickness', 'thick', 'height:', 'h:', 'h.', 't.']
    
    @classmethod
    def extract_text_from_file(cls, file_path: str) -> str:
        """Extract text from PDF or image file."""
        ext = os.path.splitext(file_path)[1].lower()
        
        if ext not in cls.SUPPORTED_EXTENSIONS:
            raise ValueError(f"Unsupported file type: {ext}")
        
        # For text files, read directly
        if ext in ['.txt', '.csv']:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()
        
        if not TESSERACT_AVAILABLE:
            print("WARNING: Tesseract OCR not available. Cannot extract text from images/PDFs.")
            return ""
        
        if ext == '.pdf':
            # First, try to extract text directly using pdfplumber (for text-based PDFs)
            text = cls._extract_text_from_pdf(file_path)
            # If no text found, fall back to OCR (for scanned/image-based PDFs)
            if not text.strip() and TESSERACT_AVAILABLE:
                print("No text extracted via pdfplumber, trying OCR...")
                text = cls._extract_from_pdf(file_path)
            return text
        elif ext in ['.doc', '.docx']:
            return cls._extract_from_word(file_path)
        elif ext in ['.xls', '.xlsx']:
            return cls._extract_from_excel(file_path)
        else:
            return cls._extract_from_image(file_path)
    
    @classmethod
    def _extract_text_from_pdf(cls, pdf_path: str) -> str:
        """Extract text directly from text-based PDF using pdfplumber."""
        if pdfplumber is None:
            print("PDFPlumber not available, skipping text extraction")
            return ""
        
        text_parts = []
        try:
            with pdfplumber.open(pdf_path) as pdf:
                print(f"PDF has {len(pdf.pages)} pages (pdfplumber)")
                for i, page in enumerate(pdf.pages):
                    text = page.extract_text()
                    if text:
                        text_parts.append(text)
                        lines = text.split('\n')
                        non_empty = [l for l in lines if l.strip()]
                        print(f"Page {i+1} (pdfplumber): {len(non_empty)} non-empty lines")
                    else:
                        print(f"Page {i+1} (pdfplumber): No text extracted")
        except Exception as e:
            print(f"PDFPlumber extraction error: {e}")
            return ""
        
        full_text = "\n".join(text_parts)
        print(f"\nPDFPlumber total text: {len(full_text)} chars")
        return full_text
    
    @classmethod
    def _extract_from_image(cls, image_path: str) -> str:
        """Extract text from image using OCR."""
        try:
            if not TESSERACT_AVAILABLE:
                return ""
            image = Image.open(image_path)
            text = pytesseract.image_to_string(image)
            print(f"Image extracted text length: {len(text)}")
            return text
        except Exception as e:
            print(f"Image extraction error: {e}")
            import traceback
            traceback.print_exc()
            return ""

    @classmethod
    def _extract_from_word(cls, file_path: str) -> str:
        """Extract text from Word document (.doc, .docx)."""
        try:
            if docx is None:
                print("python-docx not available, cannot extract Word document")
                return ""
            
            print(f"Extracting text from Word: {file_path}")
            doc = docx.Document(file_path)
            
            text_parts = []
            # Extract from paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text.strip())
            
            # Extract from tables (important for operation data)
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        text_parts.append(' '.join(row_text))
            
            full_text = '\n'.join(text_parts)
            print(f"Word document extracted: {len(full_text)} chars, {len(text_parts)} lines/rows")
            return full_text
            
        except Exception as e:
            print(f"Word extraction error: {e}")
            import traceback
            traceback.print_exc()
            return ""

    @classmethod
    def _extract_from_excel(cls, file_path: str) -> str:
        """Extract text from Excel file (.xls, .xlsx)."""
        try:
            print(f"Extracting text from Excel: {file_path}")
            text_parts = []
            
            # Try openpyxl first for .xlsx
            if file_path.endswith('.xlsx') and openpyxl is not None:
                print("Using openpyxl for .xlsx extraction")
                wb = openpyxl.load_workbook(file_path, data_only=True)
                for sheet_name in wb.sheetnames:
                    sheet = wb[sheet_name]
                    print(f"Processing sheet: {sheet_name}")
                    for row in sheet.iter_rows():
                        row_values = []
                        for cell in row:
                            if cell.value is not None:
                                row_values.append(str(cell.value).strip())
                        if row_values:
                            text_parts.append(' '.join(row_values))
                            
            # Try xlrd for .xls
            elif file_path.endswith('.xls') and xlrd is not None:
                print("Using xlrd for .xls extraction")
                wb = xlrd.open_workbook(file_path)
                for sheet_idx in range(wb.nsheets):
                    sheet = wb.sheet_by_index(sheet_idx)
                    print(f"Processing sheet: {sheet.name}")
                    for row_idx in range(sheet.nrows):
                        row_values = []
                        for col_idx in range(sheet.ncols):
                            cell_value = sheet.cell_value(row_idx, col_idx)
                            if cell_value is not None and str(cell_value).strip():
                                row_values.append(str(cell_value).strip())
                        if row_values:
                            text_parts.append(' '.join(row_values))
            else:
                print(f"No Excel library available for {file_path}")
                return ""
            
            full_text = '\n'.join(text_parts)
            print(f"Excel file extracted: {len(full_text)} chars, {len(text_parts)} rows")
            return full_text
            
        except Exception as e:
            print(f"Excel extraction error: {e}")
            import traceback
            traceback.print_exc()
            return ""
    
    @classmethod
    def extract_multiple_operations(cls, text: str) -> list:
        """Extract multiple operations from text - detects table rows and operation sections."""
        
        # First, try the new PDF format extraction (for "Operation X: Name" with two-column tables)
        pdf_format_ops = cls.extract_operations_from_pdf_format(text)
        if pdf_format_ops:
            print(f"\n✓ Extracted {len(pdf_format_ops)} operations using PDF format parser")
            return pdf_format_ops
        
        operations = []
        lines = text.split('\n')
        
        # List of valid operation types to detect in tables - EXPANDED
        valid_operations = ['turning', 'milling', 'drilling', 'grinding', 'boring', 'welding', 
                           'heat treatment', 'surface treatment', 'jig boring', 'rubber press', 
                           'jig', 'heat', 'surface']
        valid_materials = ['steel', 'aluminium', 'aluminum', 'titanium', 'brass', 'copper']
        valid_duties = ['light duty', 'medium duty', 'heavy duty', 'light', 'medium', 'heavy']
        valid_shapes = ['round', 'rectangular', 'square', 'cylindrical']
        
        print(f"\n--- Parsing {len(lines)} lines for table data ---")
        
        print(f"\n--- ALL LINES (total: {len(lines)}) ---")
        for i, line in enumerate(lines):
            stripped = line.strip()
            if stripped and len(stripped) > 3:  # Only print non-empty, meaningful lines
                print(f"  Line {i}: {repr(stripped[:150])}")
        
        # Table detection: Look for lines that start with an operation name
        # Pattern: lines containing operation type followed by material, machine, etc.
        # Also support lines like "Operation 15 Milling" or "15. Milling"
        table_rows = []
        skip_lines = set()  # Track lines to skip (merged multi-line operations)
        for i, line in enumerate(lines):
            if i in skip_lines:
                continue
            
            line_lower = line.lower().strip()
            if not line_lower:
                continue
            
            # Special case: detect "rubber" on one line followed by "press" on the next
            # (handles multi-line table entries like "Rubber Small" on line N, "press Steel..." on line N+1)
            if line_lower.startswith('rubber'):
                # Look ahead to see if next line starts with 'press'
                if i + 1 < len(lines):
                    next_line = lines[i + 1].strip().lower()
                    if next_line.startswith('press'):
                        combined = line + ' ' + lines[i + 1].strip()
                        table_rows.append({'line': combined, 'index': i, 'operation': 'rubber press'})
                        print(f"FOUND split rubber press at lines {i}-{i+1}: {combined[:80]}...")
                        skip_lines.add(i + 1)  # Skip the press line
                        continue
            
            # Check if line starts with an operation type (for table rows) - case insensitive
            matched_op = None
            for op in valid_operations:
                # Match at start of line (already lowercase)
                op_pattern = rf'^\s*{re.escape(op)}\b'
                if re.search(op_pattern, line_lower):
                    matched_op = op
                    break
                # Also try matching without word boundary for multi-word operations
                if ' ' in op:
                    simple_pattern = rf'^\s*{re.escape(op)}'
                    if re.search(simple_pattern, line_lower):
                        matched_op = op
                        break
            
            if matched_op:
                table_rows.append({'line': line, 'index': i, 'operation': matched_op})
                print(f"FOUND ROW {len(table_rows)} at line {i}: {matched_op} -> {line[:80]}...")
            else:
                # Check for "Operation X OperationType" pattern (e.g., "Operation 15 Milling")
                # or "X. OperationType" pattern (e.g., "15. Milling")
                op_prefix_patterns = [
                    rf'^\s*(?:operation|op)?\s*\d+[\.:\)\-]?\s*(turning|milling|drilling|grinding|boring|welding|heat|surface|jig|rubber)',
                    rf'^\s*\d+\.\s*(turning|milling|drilling|grinding|boring|welding|heat|surface|jig|rubber)',
                    rf'^\s*\d+\s+(turning|milling|drilling|grinding|boring|welding|heat|surface|jig|rubber)',
                ]
                for pattern in op_prefix_patterns:
                    match = re.search(pattern, line_lower)
                    if match:
                        potential_op = match.group(1).lower() if match.groups() else None
                        if not potential_op:
                            continue
                        # Check if the word after the number is a valid operation
                        for valid_op in valid_operations:
                            if valid_op in potential_op or potential_op in valid_op:
                                table_rows.append({'line': line, 'index': i, 'operation': valid_op})
                                print(f"FOUND ROW {len(table_rows)} (with prefix) at line {i}: {valid_op} -> {line[:80]}...")
                                break
                        else:
                            continue
                        break
        
        print(f"\n--- Found {len(table_rows)} potential table rows ---")
        
        # Sort table rows by their original line index to ensure PDF order is maintained
        table_rows.sort(key=lambda x: x['index'])
        print("Table rows sorted by line index (PDF order):")
        for i, row in enumerate(table_rows):
            print(f"  {i+1}. Line {row['index']}: {row['operation']}")
        
        # If we found table rows, extract each as an operation
        if len(table_rows) >= 1:
            for i, row_info in enumerate(table_rows):
                row_text = row_info['line']
                
                # For table rows, the entire data is usually in one line
                # But we might need to look at next line if it's wrapped
                if i + 1 < len(table_rows):
                    next_start = table_rows[i + 1]['index']
                    section_lines = lines[row_info['index']:next_start]
                else:
                    # Last row - take rest of document
                    section_lines = lines[row_info['index']:]
                
                section_text = '\n'.join(section_lines)
                print(f"\n--- Processing table row {i+1}: {row_info['operation']} ---")
                
                # Extract details from this row
                details = cls._extract_operation_from_line(section_text, row_info['operation'])
                
                # Normalize operation type to use frontend-compatible values
                if details['operation_type']:
                    op_type = details['operation_type'].lower().strip()
                    # Map to frontend-supported operation types
                    if 'jig' in op_type or op_type in ['jig', 'jig boring']:
                        op_type = 'boring'  # Frontend expects 'boring', not 'jig_boring'
                    elif 'rubber' in op_type or op_type in ['rubber', 'rubber press']:
                        op_type = 'rubber_press'
                    elif ' ' in op_type:
                        op_type = op_type.replace(' ', '_')
                    details['operation_type'] = op_type
                
                if any(v is not None for v in details.values()):
                    operations.append(details)
                    print(f"Row {i+1} extracted: {details}")
                else:
                    print(f"Row {i+1}: No valid data")
        
            # If no table rows found via line-matching, try to find operations that span multiple lines
            # by looking for operation keywords followed by continuation lines
            if not table_rows:
                print("\n--- Trying multi-line operation detection ---")
                i = 0
                while i < len(lines):
                    line = lines[i].strip()
                    line_lower = line.lower()
                    if not line_lower:
                        i += 1
                        continue
                    
                    # Check if this line starts with an operation
                    for op in valid_operations:
                        if re.search(rf'^\s*{re.escape(op)}\b', line_lower):
                            # Found an operation start - look ahead for continuation
                            combined = line
                            j = i + 1
                            while j < len(lines) and j < i + 3:  # Check next 2 lines
                                next_line = lines[j].strip()
                                next_lower = next_line.lower()
                                # If next line starts with a continuation word, merge it
                                if next_lower.startswith('press') or next_lower.startswith('treatment'):
                                    combined += ' ' + next_line
                                    print(f"  Merged continuation at line {j}: {next_line[:60]}")
                                    j += 1
                                else:
                                    break
                            table_rows.append({'line': combined, 'index': i, 'operation': op})
                            print(f"FOUND multi-line operation at {i}: {op}")
                            i = j  # Skip the merged lines
                            break
                    else:
                        i += 1
        
        return operations
    
    @classmethod
    def _extract_operation_from_line(cls, text: str, detected_operation: str = None) -> dict:
        """Extract operation details from a table row or line of text."""
        details = {
            'operation_type': detected_operation,
            'material': None,
            'machine': None,
            'man_hours': None,
            'duty_category': None,
            'diameter': None,
            'length': None,
            'breadth': None,
            'height': None,
            'shape': None
        }
        
        text_lower = text.lower()
        lines = text.split('\n')
        
        # For single line (table row), extract all values from that line
        if len(lines) == 1 or (len(lines) > 0 and all(len(l.strip()) < 100 for l in lines[:2])):
            main_line = lines[0]
            main_lower = main_line.lower()
            
            # Extract material
            materials = ['steel', 'aluminium', 'aluminum', 'titanium', 'brass', 'copper']
            for mat in materials:
                if mat in main_lower:
                    details['material'] = mat if mat != 'aluminum' else 'aluminium'
                    break
            
            # Extract duty category - improved detection with debug
            print(f"  Looking for duty in: {main_lower[:100]}...")
            duties = [
                ('heavy duty', 'heavy'), 
                ('medium duty', 'medium'), 
                ('light duty', 'light'),
                ('heavy', 'heavy'),  # Also check standalone
                ('medium', 'medium'),
                ('light', 'light')
            ]
            duty_found = False
            for duty_full, duty_short in duties:
                if duty_full in main_lower:
                    details['duty_category'] = duty_short
                    print(f"  ✓ Found duty: {duty_short} (matched '{duty_full}')")
                    duty_found = True
                    break
            
            if not duty_found:
                print(f"  ✗ No duty found in line - checking nearby lines...")
                # Look for duty in the full section text (multiple lines)
                full_lower = text_lower
                for duty_full, duty_short in duties:
                    if duty_full in full_lower:
                        details['duty_category'] = duty_short
                        print(f"  ✓ Found duty in section: {duty_short} (matched '{duty_full}')")
                        duty_found = True
                        break
            
            # Extract shape
            shapes = ['round', 'rectangular', 'square', 'cube', 'cylindrical']
            for shape in shapes:
                if shape in main_lower:
                    details['shape'] = shape
                    break
            
            # Extract machine name directly from the line
            # Common machine patterns - expanded
            # PRIORITY ORDER: Specific machines first, then generic patterns
            machine_patterns = [
                # CNC patterns FIRST (before standalone CNC)
                (r'\b(?:3|4|5)\s*axis\s+cnc\b', None),
                (r'turning\s+cnc', 'Turning CNC'),
                (r'cnc\s+(?:lathe|milling|machine)?', None),
                (r'cnc\b', 'CNC'),
                # Rubber press sizes - check these BEFORE generic rubber press
                (r'small\s*\(?\s*less\s+than\s*650', 'Small (less than 650 MMX 650 MM)'),
                (r'medium\s*\(?\s*more\s+than\s*650', 'Medium (more than 650 MMX 650 MM)'),
                (r'small\s*\([^)]*650[^)]*\)', 'Small (less than 650 MMX 650 MM)'),
                (r'medium\s*\([^)]*650[^)]*\)', 'Medium (more than 650 MMX 650 MM)'),
                (r'large\s*\([^)]*\)', 'Large'),
                (r'she[l1]\b', 'Small (less than 650 MMX 650 MM)'),  # OCR error: Shel -> Small
                (r'med[i1]um\b', 'Medium (more than 650 MMX 650 MM)'),  # OCR error: med1um -> Medium
                (r'small\b', 'Small (less than 650 MMX 650 MM)'),
                (r'medium\b', 'Medium (more than 650 MMX 650 MM)'),
                (r'large\b', 'Large'),
                # Generic rubber press (only if size not found)
                (r'rubber\s*press', 'Rubber Press'),
                # Other machines
                (r'(?:drilling|milling|grinding|turning|boring)\s+machine', None),
                (r'\bspm\b', 'SPM'),
                (r'jig\s+boring', 'Jig boring'),
                (r'conventional', 'Conventional'),
                (r'turning', None),
                (r'lathe', None),
                (r'drilling', None),
                (r'grinding', None),
                (r'milling', None),
                (r'special\s+purpose\s*machine(?:\s*\([^)]*\))?','Special Purpose Machine(SPM)'),
            ]
            for pattern, default_name in machine_patterns:
                match = re.search(pattern, main_lower, re.IGNORECASE)
                if match:
                    # Get the actual matched text with original casing
                    matched_text = main_line[match.start():match.end()]
                    # Normalize casing
                    if default_name:
                        details['machine'] = default_name
                    else:
                        # Capitalize first letter of each word
                        details['machine'] = matched_text.strip().title()
                    print(f"  Found machine: {details['machine']}")
                    break
            
            # Extract all numeric values from the line
            # First, remove common prefixes that contain operation numbers
            # Pattern: "Operation 12" or "Op 12" at the start
            cleaned_line = re.sub(r'^\s*(?:operation|op)\s*\d+\s*', ' ', main_line, flags=re.IGNORECASE)
            # Remove machine size descriptors that contain numbers (like "650 MMX")
            cleaned_line = re.sub(r'\d+\s*mmx?\s*\d*\s*mm\)?', ' ', cleaned_line, flags=re.IGNORECASE)
            cleaned_line = re.sub(r'less\s+than\s+\d+', ' ', cleaned_line, flags=re.IGNORECASE)
            cleaned_line = re.sub(r'more\s+than\s+\d+', ' ', cleaned_line, flags=re.IGNORECASE)
            # Remove standalone numbers that are likely part of machine descriptions (650, 650MM, etc)
            cleaned_line = re.sub(r'\b65[0-9]\b', ' ', cleaned_line, flags=re.IGNORECASE)
            numbers = re.findall(r'\b(\d+(?:\.\d+)?)\b', cleaned_line)
            numbers = [float(n) for n in numbers]
            print(f"  Found numbers (after cleaning): {numbers}")
            
            # Special handling for rubber press operations
            is_rubber_press = detected_operation and 'rubber' in detected_operation.lower()
            if is_rubber_press:
                # Extract machine name for rubber press
                if details['machine'] is None:
                    details['machine'] = cls._extract_machine(text, lines)
                    print(f"  Extracted rubber press machine: {details['machine']}")
                
                print(f"  Rubber press detected - processing {len(numbers)} numbers")
                # Rubber press format: Hours, Diameter, Length, Breadth, Height (no setup/cycle time)
                if len(numbers) >= 5:
                    details['man_hours'] = numbers[0]
                    details['diameter'] = numbers[1]
                    details['length'] = numbers[2]
                    details['breadth'] = numbers[3] if numbers[3] > 0 else None
                    details['height'] = numbers[4] if numbers[4] > 0 else None
                    print(f"  Rubber press 5-col: Hours={numbers[0]}, Dia={numbers[1]}, Len={numbers[2]}, Breadth={numbers[3]}, Height={numbers[4]}")
                elif len(numbers) >= 4:
                    details['man_hours'] = numbers[0]
                    details['diameter'] = numbers[1]
                    details['length'] = numbers[2]
                    details['breadth'] = numbers[3] if numbers[3] > 0 else None
                    print(f"  Rubber press 4-col: Hours={numbers[0]}, Dia={numbers[1]}, Len={numbers[2]}, Breadth={numbers[3]}")
                elif len(numbers) >= 3:
                    details['man_hours'] = numbers[0]
                    details['diameter'] = numbers[1]
                    details['length'] = numbers[2]
                    print(f"  Rubber press 3-col: Hours={numbers[0]}, Dia={numbers[1]}, Len={numbers[2]}")
            else:
                # Standard handling for non-rubber operations
                # Map numbers to fields based on position and context
                # Table order typically: [Hours, Diameter, Length, Breadth, Height]
                # OR for rectangular parts (milling): [Hours, -, Length, Breadth, Height]
                # OR for round parts: [Hours, Diameter, Length, -, -]
                
                # Based on typical table structure, try to identify each value by position
                # Column indices: 0=Hours, 1=Diameter/Dash, 2=Length, 3=Breadth/Dash, 4=Height/Dash
                if len(numbers) >= 5:
                    # Full 5-column table: Hours, Diameter, Length, Breadth, Height
                    details['man_hours'] = numbers[0]  # Large: 1200-4200
                    # Index 1 could be diameter or dash (represented as 0 or just skipped)
                    if numbers[1] > 0:
                        details['diameter'] = numbers[1]
                    details['length'] = numbers[2]
                    if numbers[3] > 0:
                        details['breadth'] = numbers[3]
                    if numbers[4] > 0:
                        details['height'] = numbers[4]
                elif len(numbers) >= 4:
                    # 4-column table: Hours, Dim1, Dim2, Dim3
                    # ONLY milling gets length/breadth/height (rectangular)
                    # ALL other operations get diameter/length (round) - shape doesn't matter
                    is_milling = detected_operation and 'mill' in detected_operation.lower()
                    
                    details['man_hours'] = numbers[0]
                    
                    if is_milling:
                        # Milling rectangular: Length, Breadth, Height
                        details['length'] = numbers[1]
                        details['breadth'] = numbers[2]
                        details['height'] = numbers[3]
                        print(f"  4-column MILLING: Hours={numbers[0]}, Length={numbers[1]}, Breadth={numbers[2]}, Height={numbers[3]}")
                    else:
                        # ALL other operations: Diameter, Length (round)
                        details['diameter'] = numbers[1]
                        details['length'] = numbers[2]
                        print(f"  4-column NON-MILLING: Hours={numbers[0]}, Diameter={numbers[1]}, Length={numbers[2]}")
                elif len(numbers) >= 3:
                    # Standard 3-column table: Hours, Diameter, Length
                    # For round parts (most operations except milling): Diameter is at index 1
                    # For rectangular parts: Length is at index 1, Breadth at index 2
                    details['man_hours'] = numbers[0]
                    # ALWAYS set diameter and length for 3-column tables
                    # Position 1 is always diameter for round parts
                    # Position 2 is always length
                    details['diameter'] = numbers[1]
                    details['length'] = numbers[2]
                    print(f"  3-col ROUND: Hours={numbers[0]}, Dia={numbers[1]}, Len={numbers[2]}")
                elif len(numbers) >= 2:
                    # 2 columns - could be without several fields
                    # Try to identify based on value ranges
                    for num in numbers:
                        if num > 1000 and details['man_hours'] is None:
                            details['man_hours'] = num
                        elif 50 < num < 1000:
                            # Could be length or diameter
                            if details['length'] is None:
                                details['length'] = num
                            elif details['diameter'] is None:
                                details['diameter'] = num
                        elif num > 10 and details['diameter'] is None:
                            details['diameter'] = num
                        elif details['length'] is None:
                            details['length'] = num
                else:
                    # Less than 2 numbers - use range-based logic
                    for num in numbers:
                        if num > 1000:  # Work hours
                            if details['man_hours'] is None:
                                details['man_hours'] = num
                        elif 100 < num <= 1000:  # Length or diameter
                            if details['length'] is None:
                                details['length'] = num
                            elif details['diameter'] is None:
                                details['diameter'] = num
                        elif 20 <= num <= 100:  # Diameter
                            if details['diameter'] is None:
                                details['diameter'] = num
                            elif details['length'] is None:
                                details['length'] = num
                        elif num > 0:  # Small numbers
                            if details['diameter'] is None:
                                details['diameter'] = num
                            elif details['length'] is None:
                                details['length'] = num
        
        # Use the regular extraction as fallback/supplement
        regular = cls.extract_operation_details(text)
        for key in details:
            if details[key] is None and regular.get(key) is not None:
                details[key] = regular[key]
        
        # Ensure operation type is set
        if details['operation_type'] is None:
            details['operation_type'] = regular.get('operation_type')
        
        return details
    
    @classmethod
    def _extract_sections(cls, text: str) -> list:
        """Extract operations based on section boundaries (Operation 1, Op 2, etc.)."""
        operations = []
        
        # Patterns that indicate operation boundaries
        operation_patterns = [
            r'(?:^|\n)\s*(?:operation|op|opn)\s*(\d+)[\s:\.\-]',
            r'(?:^|\n)\s*(\d+)\s*[\.\)\-]\s*(?:operation|op)',
        ]
        
        boundaries = []
        for pattern in operation_patterns:
            for match in re.finditer(pattern, text, re.IGNORECASE):
                op_num = match.group(1) if match.groups() else None
                boundaries.append({
                    'start': match.start(),
                    'number': int(op_num) if op_num else None,
                    'matched_text': match.group(0)
                })
        
        boundaries = sorted(boundaries, key=lambda x: x['start'])
        
        # Remove duplicates that are too close
        unique_boundaries = []
        last_start = -100
        for b in boundaries:
            if b['start'] - last_start > 50:
                unique_boundaries.append(b)
                last_start = b['start']
        boundaries = unique_boundaries
        
        print(f"Found {len(boundaries)} section boundaries")
        
        if len(boundaries) <= 1:
            details = cls.extract_operation_details(text)
            if any(v is not None for v in details.values()):
                return [details]
            return []
        
        # Extract each section
        for i, boundary in enumerate(boundaries):
            section_start = boundary['start']
            if i + 1 < len(boundaries):
                section_end = boundaries[i + 1]['start']
            else:
                section_end = len(text)
            
            section_text = text[section_start:section_end]
            details = cls.extract_operation_details(section_text)
            
            if any(v is not None for v in details.values()):
                operations.append(details)
        
        return operations
    
    @classmethod
    def extract_operation_details(cls, text: str) -> dict:
        """Extract operation details from text."""
        details = {
            'operation_type': None,
            'material': None,
            'machine': None,
            'man_hours': None,
            'duty_category': None,
            'diameter': None,
            'length': None,
            'breadth': None,
            'height': None,
            'shape': None
        }
        
        text_lower = text.lower()
        lines = text.split('\n')
        
        # Extract each field
        details['operation_type'] = cls._extract_value(text, lines, cls.OPERATION_KEYWORDS, 
                                                      ['turning', 'milling', 'drilling', 'grinding', 'boring'])
        details['material'] = cls._extract_value(text, lines, cls.MATERIAL_KEYWORDS,
                                               ['steel', 'aluminium', 'aluminum', 'titanium', 'brass', 'copper'])
        details['machine'] = cls._extract_machine(text, lines)
        details['man_hours'] = cls._extract_numeric_value(text, lines, cls.MAN_HOURS_KEYWORDS)
        details['duty_category'] = cls._extract_value(text, lines, cls.DUTY_KEYWORDS,
                                                      ['light', 'medium', 'heavy'])
        details['diameter'] = cls._extract_numeric_value(text, lines, cls.DIAMETER_KEYWORDS)
        details['length'] = cls._extract_numeric_value(text, lines, cls.LENGTH_KEYWORDS)
        details['breadth'] = cls._extract_numeric_value(text, lines, cls.BREADTH_KEYWORDS)
        details['height'] = cls._extract_numeric_value(text, lines, cls.HEIGHT_KEYWORDS)
        
        # Determine shape based on dimensions found
        if details['diameter'] and not details['breadth']:
            details['shape'] = 'round'
        elif details['breadth'] or details['height']:
            details['shape'] = 'rectangular'
        
        return details
    
    @classmethod
    def _extract_value(cls, text: str, lines: list, keywords: list, allowed_values: list) -> Optional[str]:
        """Extract a value based on keywords and allowed values - case insensitive."""
        text_lower = text.lower()
        
        for i, line in enumerate(lines):
            line_lower = line.lower()
            # Check if any keyword is in the line
            if any(kw.lower() in line_lower for kw in keywords):
                # Look for allowed values in the same line or next few lines (case insensitive)
                search_text = line_lower
                if i + 1 < len(lines):
                    search_text += " " + lines[i + 1].lower()
                
                for value in allowed_values:
                    if value.lower() in search_text:
                        # Normalize aluminium spelling
                        if value.lower() in ['aluminium', 'aluminum']:
                            return 'aluminium'
        if re.search(r'\bcnc\b', text_lower):
            # Make sure it's not already part of a longer machine name we extracted
            return 'CNC'
        
        for line in lines:
            line_lower = line.lower()
            if any(kw.lower() in line_lower for kw in cls.MACHINE_KEYWORDS):
                # For rubber press lines, check for size patterns first
                if 'rubber' in line_lower:
                    for pattern, default_name in rubber_press_patterns:
                        match = re.search(pattern, line_lower, re.IGNORECASE)
                        if match:
                            return default_name
                
                # Common machine patterns - expanded
                machine_patterns = [
                    r'\b(?:3|4|5)\s*axis\s+cnc\b',  # Check 3/4/5 axis CNC FIRST
                    r'cnc\s+(?:lathe|milling|machine)',  # CNC with suffix
                    r'lathe', r'milling', r'drilling', r'grinding',
                    r'jig\s*boring', r'vmc', r'hmc', r'turning\s+cnc', r'turning',
                    r'rubber\s*press', r'spm',
                    r'conventional'
                ]
                for pattern in machine_patterns:
                    match = re.search(pattern, line_lower, re.IGNORECASE)
                    if match:
                        # Get the matched text and expand to capture full machine name
                        words = line.split()
                        matched_word_idx = None
                        for wi, word in enumerate(words):
                            if re.search(pattern, word, re.IGNORECASE):
                                matched_word_idx = wi
                                break
                        
                        if matched_word_idx is not None:
                            # Get up to 2 words before and 3 words after for longer names
                            start_idx = max(0, matched_word_idx - 2)
                            end_idx = min(len(words), matched_word_idx + 4)
                            machine_name = ' '.join(words[start_idx:end_idx]).strip()
                            return machine_name
                        
                        # Fallback: return matched text
                        return match.group(0)
        return None
    
    @classmethod
    def _extract_numeric_value(cls, text: str, lines: list, keywords: list) -> Optional[float]:
        """Extract a numeric value based on keywords - case insensitive."""
        for i, line in enumerate(lines):
            line_lower = line.lower()
            if any(kw.lower() in line_lower for kw in keywords):
                # Look for number patterns in this line and next line
                search_text = line
                if i + 1 < len(lines):
                    search_text += " " + lines[i + 1]
                
                # Find numbers (including decimals)
                numbers = re.findall(r'\d+\.?\d*', search_text)
                if numbers:
                    try:
                        return float(numbers[0])
                    except ValueError:
                        continue
        return None

    @classmethod
    def extract_operations_from_pdf_format(cls, text: str) -> list:
        """Extract operations from PDF with 'Operation X: Name' header format and two-column tables."""
        operations = []
        lines = text.split('\n')
        
        print(f"\n=== PDF Format Extraction: {len(lines)} lines ===")
        
        # Pattern to match "Operation X: Name" - more flexible
        operation_header_pattern = re.compile(
            r'operation\s*(\d+)\s*[:\-]?\s*(turning|milling|drilling|grinding|boring|welding|heat treatment|surface treatment|jig boring|rubber press|rubber)',
            re.IGNORECASE
        )
        
        # Alternative: just match "Operation X" and capture next words
        simple_op_pattern = re.compile(
            r'operation\s*(\d+)',
            re.IGNORECASE
        )
        
        i = 0
        while i < len(lines):
            line = lines[i].strip()
            line_lower = line.lower()
            
            if not line:
                i += 1
                continue
            
            # Check for operation header with operation type
            op_match = operation_header_pattern.search(line)
            
            # If not found, try simple pattern and look for operation type in same line
            if not op_match:
                simple_match = simple_op_pattern.search(line)
                if simple_match:
                    op_num = simple_match.group(1)
                    # Look for operation type in the rest of the line
                    op_type = None
                    for opt in ['turning', 'milling', 'drilling', 'grinding', 'boring', 'welding', 'heat', 'surface', 'jig', 'rubber']:
                        if opt in line_lower:
                            op_type = opt
                            break
                    if op_type:
                        op_match = simple_match
            
            if op_match:
                op_num = op_match.group(1)
                # Try to get operation type from the match, otherwise detect from line
                try:
                    op_type = op_match.group(2).lower() if hasattr(op_match, 'group') and len(op_match.groups()) > 1 else None
                except:
                    op_type = None
                
                if not op_type:
                    # Detect from the line
                    for opt in ['turning', 'milling', 'drilling', 'grinding', 'boring', 'welding', 'heat', 'surface', 'jig', 'rubber']:
                        if opt in line_lower:
                            op_type = opt
                            break
                
                # Normalize operation type
                if op_type:
                    if 'jig' in op_type:
                        op_type = 'boring'
                    elif 'rubber' in op_type:
                        op_type = 'rubber_press'
                    elif 'heat' in op_type:
                        op_type = 'heat_treatment'
                    elif 'surface' in op_type:
                        op_type = 'surface_treatment'
                    elif 'mill' in op_type:
                        op_type = 'milling'
                    elif ' ' in op_type:
                        op_type = op_type.replace(' ', '_')
                else:
                    op_type = 'unknown'
                
                print(f"\n--- Found Operation {op_num}: {op_type} at line {i} ---")
                print(f"  Line content: {line[:100]}")
                
                # Collect all lines until next operation header or end
                operation_lines = []
                j = i + 1
                while j < len(lines):
                    next_line = lines[j].strip()
                    # Check if next line is a new operation header
                    if simple_op_pattern.search(next_line.lower()):
                        break
                    if next_line:
                        operation_lines.append(next_line)
                    j += 1
                
                print(f"  Collected {len(operation_lines)} lines for operation")
                
                # Parse operation details from collected lines
                details = cls._parse_generated_pdf_table('\n'.join(operation_lines), op_type)
                
                if details:
                    operations.append(details)
                    print(f"  ✓ Extracted: {details}")
                else:
                    print(f"  ✗ Failed to parse operation details")
                
                i = j  # Skip to next operation
            else:
                i += 1
        
        print(f"\n=== Total operations extracted: {len(operations)} ===")
        return operations

    @classmethod
    def _parse_generated_pdf_table(cls, text: str, operation_type: str) -> Optional[dict]:
        """Parse generated PDF table format with two-column Label: Value pairs."""
        lines = text.split('\n')
        
        details = {
            'operation_type': operation_type,
            'material': None,
            'machine': None,
            'man_hours': None,
            'duty_category': None,
            'diameter': None,
            'length': None,
            'breadth': None,
            'height': None,
            'shape': None
        }
        
        # Join all lines and also keep individual lines for pattern matching
        full_text = ' '.join(lines)
        full_text_lower = full_text.lower()
        
        print(f"  Parsing PDF table: {len(lines)} lines")
        print(f"  Full text preview: {full_text[:200]}...")
        
        # Extract Material - look for "Material:" followed by value
        material_patterns = [
            r'material[:\s]+(\w+)',
            r'material\s+(\w+)',
        ]
        for pattern in material_patterns:
            match = re.search(pattern, full_text_lower)
            if match:
                mat = match.group(1).lower()
                if mat in ['steel', 'aluminium', 'aluminum', 'titanium', 'brass', 'copper']:
                    details['material'] = mat if mat != 'aluminum' else 'aluminium'
                    print(f"    Found material: {details['material']}")
                    break
        
        # Extract Machine - look for "Machine:" followed by machine name
        machine_patterns = [
            r'machine[:\s]+([^\n]+?)(?:\s+(?:duty|setup|cycle|man|diameter|length|shape|$))',
            r'machine[:\s]+([^\n,]+)',
        ]
        for pattern in machine_patterns:
            match = re.search(pattern, full_text_lower)
            if match:
                machine = match.group(1).strip()
                if machine and machine not in ['', 'duty', 'setup', 'cycle']:
                    details['machine'] = machine
                    print(f"    Found machine: {details['machine']}")
                    break
        
        # If pattern didn't work, try simple search
        if not details['machine']:
            # Look for common machine names in the text
            machine_keywords = ['drilling machine', 'cnc', 'lathe', 'milling machine', '3 axis cnc', '5 axis cnc', 
                              'grinding machine', 'welding machine', 'boring machine', 'press', 'heat treatment']
            for keyword in machine_keywords:
                if keyword in full_text_lower:
                    details['machine'] = keyword.title()
                    print(f"    Found machine (keyword): {details['machine']}")
                    break
        
        # Extract Man Hours/Unit - handle "Man Hours/Unit:" format
        man_hours_patterns = [
            r'man\s+hours?[/\s]+unit[:\s]+(\d+(?:\.\d+)?)',
            r'man\s+hours?[:\s]+(\d+(?:\.\d+)?)',
            r'hours?[/\s]+unit[:\s]+(\d+(?:\.\d+)?)',
            r'man\s+hrs?[:\s]+(\d+(?:\.\d+)?)',
        ]
        for pattern in man_hours_patterns:
            match = re.search(pattern, full_text_lower)
            if match:
                details['man_hours'] = float(match.group(1))
                print(f"    Found man_hours: {details['man_hours']}")
                break
        
        # Extract Duty Category
        duty_patterns = [
            r'duty\s+category[:\s]+(light|medium|heavy)',
            r'duty[:\s]+(light|medium|heavy)',
        ]
        for pattern in duty_patterns:
            match = re.search(pattern, full_text_lower)
            if match:
                details['duty_category'] = match.group(1).lower()
                print(f"    Found duty_category: {details['duty_category']}")
                break
        
        # If no duty category found, search in full text
        if not details['duty_category']:
            for duty in ['heavy duty', 'medium duty', 'light duty']:
                if duty in full_text_lower:
                    details['duty_category'] = duty.split()[0]
                    print(f"    Found duty_category (text): {details['duty_category']}")
                    break
            # Also check for standalone duty words
            if not details['duty_category']:
                for duty in ['heavy', 'medium', 'light']:
                    pattern = rf'\b{duty}\b'
                    if re.search(pattern, full_text_lower):
                        # Make sure it's not part of machine name or material
                        if duty not in ['material', 'machine']:
                            details['duty_category'] = duty
                            print(f"    Found duty_category (word): {details['duty_category']}")
                            break
        
        # Extract Setup Time (min)
        setup_patterns = [
            r'setup\s+time\s*\(?min\)?[:\s]+(\d+(?:\.\d+)?)',
            r'setup\s+time[:\s]+(\d+(?:\.\d+)?)',
            r'setup[:\s]+(\d+(?:\.\d+)?)',
        ]
        for pattern in setup_patterns:
            match = re.search(pattern, full_text_lower)
            if match:
                details['machine_setup_time'] = float(match.group(1))
                print(f"    Found setup_time: {details['machine_setup_time']}")
                break
        
        # Extract Cycle Time (min)
        cycle_patterns = [
            r'cycle\s+time\s*\(?min\)?[:\s]+(\d+(?:\.\d+)?)',
            r'cycle\s+time[:\s]+(\d+(?:\.\d+)?)',
            r'cycle[:\s]+(\d+(?:\.\d+)?)',
        ]
        for pattern in cycle_patterns:
            match = re.search(pattern, full_text_lower)
            if match:
                details['cycle_time'] = float(match.group(1))
                print(f"    Found cycle_time: {details['cycle_time']}")
                break
        
        # Extract dimensions based on operation type
        is_milling = 'mill' in operation_type.lower()
        
        if is_milling:
            # For milling: Length, Breadth, Height (no diameter)
            length_match = re.search(r'length[:\s]+(\d+(?:\.\d+)?)', full_text_lower)
            if length_match:
                details['length'] = float(length_match.group(1))
                print(f"    Found length: {details['length']}")
            
            breadth_match = re.search(r'breadth[:\s]+(\d+(?:\.\d+)?)', full_text_lower)
            if breadth_match:
                details['breadth'] = float(breadth_match.group(1))
                print(f"    Found breadth: {details['breadth']}")
            
            height_match = re.search(r'height[:\s]+(\d+(?:\.\d+)?)', full_text_lower)
            if height_match:
                details['height'] = float(height_match.group(1))
                print(f"    Found height: {details['height']}")
        else:
            # For other operations: Diameter, Length
            diameter_match = re.search(r'diameter[:\s]+(\d+(?:\.\d+)?)', full_text_lower)
            if diameter_match:
                details['diameter'] = float(diameter_match.group(1))
                print(f"    Found diameter: {details['diameter']}")
            
            length_match = re.search(r'length[:\s]+(\d+(?:\.\d+)?)', full_text_lower)
            if length_match:
                details['length'] = float(length_match.group(1))
                print(f"    Found length: {details['length']}")
        
        # Extract Shape
        shape_patterns = [
            r'shape[:\s]+(round|rectangular|square|cylindrical)',
            r'shape\s+(round|rectangular|square|cylindrical)',
        ]
        for pattern in shape_patterns:
            match = re.search(pattern, full_text_lower)
            if match:
                details['shape'] = match.group(1).lower()
                print(f"    Found shape: {details['shape']}")
                break
        
        # If shape not found via pattern, search in text
        if not details['shape']:
            for shape in ['round', 'rectangular', 'square', 'cylindrical']:
                if shape in full_text_lower:
                    details['shape'] = shape
                    print(f"    Found shape (text): {details['shape']}")
                    break
        
        # Debug output
        print(f"  Parsed details: material={details['material']}, machine={details['machine']}, "
              f"man_hours={details['man_hours']}, duty={details['duty_category']}")
        
        # If we have at least material or machine or man_hours, consider it valid
        if any([details['material'], details['machine'], details['man_hours']]):
            return details
        
        print(f"  ✗ Not enough data extracted - rejecting operation")
        return None


def extract_from_uploaded_file(file_path: str) -> dict:
    """Extract operation details from an uploaded file."""
    print(f"\n{'='*60}")
    print(f"EXTRACTING FROM FILE: {file_path}")
    print(f"File exists: {os.path.exists(file_path)}")
    print(f"File size: {os.path.getsize(file_path) if os.path.exists(file_path) else 'N/A'} bytes")
    print(f"PDFPlumber available: {pdfplumber is not None}")
    print(f"Tesseract available: {TESSERACT_AVAILABLE}")
    print(f"{'='*60}\n")
    
    result = {
        "text_preview": "",
        "extracted_data": {},
        "success": False,  # Changed from extraction_success to success
        "operations": []
    }
    
    try:
        # Extract text from file
        text = FileExtractionService.extract_text_from_file(file_path)
        result["text_preview"] = text[:500] if text else ""
        
        print(f"\n--- Extracted text length: {len(text)} chars ---")
        print(f"First 200 chars: {text[:200]}")
        print(f"Last 200 chars: {text[-200:] if len(text) > 200 else text}")
        
        if not text or not text.strip():
            print("ERROR: No text extracted from file!")
            result["error"] = "No text could be extracted from the file"
            return result
        
        # Try to extract operations
        operations = FileExtractionService.extract_multiple_operations(text)
        print(f"\n--- Extracted {len(operations)} operations ---")
        for i, op in enumerate(operations):
            print(f"  Op {i+1}: {op.get('operation_type', 'unknown')} - {op.get('material', 'no material')}")
        
        result["operations"] = operations
        result["extracted_data"] = {"operation_count": len(operations), "operations": operations}
        result["success"] = len(operations) > 0  # Changed from extraction_success to success
        
    except Exception as e:
        print(f"ERROR during extraction: {e}")
        import traceback
        traceback.print_exc()
        result["error"] = str(e)
    
    print(f"\n--- Final result: success={result['success']}, operations={len(result.get('operations', []))}")
    return result
